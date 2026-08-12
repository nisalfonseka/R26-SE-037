from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "fullSummary.md"
OUTPUT = ROOT / "FullSummary.docx"

INK = "172A3A"
BLUE = "1F4E79"
BLUE_DARK = "17365D"
BLUE_LIGHT = "E8EEF5"
LIGHT = "F7F9FC"
MUTED = "5D6874"
GRID = "B8C4D0"
WHITE = "FFFFFF"
CAUTION = "FFF4CE"
RED = "9B1C1C"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
MARGIN_DXA = 1440
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_borders(table, color=GRID, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths: list[int]) -> None:
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border_bottom(paragraph, color=BLUE, size="10", space="5") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def add_field(paragraph, instruction: str, display: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=9, color=MUTED)


def clean_inline(text: str) -> str:
    return text.replace("\\|", "|").replace("  ", " ")


INLINE_RE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")


def add_inline(paragraph, text: str, *, base_size=10.25, color=INK, bold=False) -> None:
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(clean_inline(text[pos:match.start()]))
            set_font(run, size=base_size, color=color, bold=bold)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, name="Menlo", size=max(7.0, base_size - 1.0), color=BLUE_DARK, bold=bold)
            run.font.highlight_color = None
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=base_size, color=color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=base_size, color=color, bold=bold, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(clean_inline(text[pos:]))
        set_font(run, size=base_size, color=color, bold=bold)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.25)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, BLUE_DARK, 10, 5),
        "Heading 4": (10.5, BLUE_DARK, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True
    styles["Heading 1"].paragraph_format.page_break_before = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.25)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.widow_control = True

    code = styles.add_style("Code Block", 1)
    code.font.name = "Menlo"
    code.font.size = Pt(6.3)
    code.font.color.rgb = RGBColor.from_string(BLUE_DARK)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(5)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.line_spacing = 0.9
    code.paragraph_format.keep_together = True

    note = styles.add_style("Evidence Note", 1)
    note.font.name = "Calibri"
    note.font.size = Pt(9.25)
    note.font.color.rgb = RGBColor.from_string(BLUE_DARK)
    note.paragraph_format.left_indent = Inches(0.18)
    note.paragraph_format.right_indent = Inches(0.18)
    note.paragraph_format.space_before = Pt(5)
    note.paragraph_format.space_after = Pt(7)
    note.paragraph_format.line_spacing = 1.15


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("SinhalaJournal-LLM / SINAI  |  Repository Research Audit")
    set_font(run, size=8.5, color=MUTED, bold=True)
    set_paragraph_border_bottom(p, color=GRID, size="5", space="3")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(2)
    lead = p.add_run("R26-SE-037   |   Page ")
    set_font(lead, size=8.5, color=MUTED)
    add_field(p, " PAGE ", "1")
    tail = p.add_run(" of ")
    set_font(tail, size=8.5, color=MUTED)
    add_field(p, " NUMPAGES ", "1")


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(78)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RESEARCH IMPLEMENTATION AUDIT")
    set_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("SinhalaJournal-LLM / SINAI")
    set_font(r, size=28, color=BLUE_DARK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("Complete Repository Research Audit")
    set_font(r, size=16, color=MUTED)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(24)
    set_paragraph_border_bottom(rule, color=BLUE, size="14", space="5")

    metadata = [
        ("Project ID", "R26-SE-037"),
        ("Snapshot", "12 August 2026"),
        ("Coverage", "Training, datasets, evaluation, architecture, applications, deployment, and Git history"),
        ("Evidence rule", "Experiment artifacts and executable source supersede stale README claims"),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (label, value) in enumerate(metadata):
        for cell in table.rows[i].cells:
            set_cell_margins(cell, 95, 120, 95, 120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(table.cell(i, 0), BLUE_LIGHT)
        p1 = table.cell(i, 0).paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        add_inline(p1, label, base_size=9.5, color=BLUE_DARK, bold=True)
        p2 = table.cell(i, 1).paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        add_inline(p2, value, base_size=9.5, color=INK)
    set_table_geometry(table, [1800, 7560])
    set_table_borders(table, color=GRID, size="5")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Authoritative factual snapshot for academic paper preparation")
    set_font(r, size=10.5, color=MUTED, italic=True)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Contents")
    set_font(r, size=20, color=BLUE_DARK, bold=True)
    set_paragraph_border_bottom(p, color=BLUE, size="10", space="4")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("In Word, right-click the contents list and choose Update Field if page numbers do not refresh automatically.")
    set_font(r, size=9, color=MUTED, italic=True)

    toc_p = doc.add_paragraph()
    add_field(toc_p, ' TOC \\o "1-3" \\h \\z \\u ', "Table of contents")
    doc.add_page_break()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        raw = lines[i].strip()
        cells = [c.strip() for c in raw.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def compute_widths(rows: list[list[str]]) -> list[int]:
    cols = max(len(r) for r in rows)
    scores = []
    for c in range(cols):
        values = [len(re.sub(r"[`*]", "", r[c])) if c < len(r) else 0 for r in rows]
        max_len = max(values or [1])
        avg_len = sum(values) / max(1, len(values))
        score = max(7.0, min(45.0, max_len * 0.42 + avg_len * 0.58))
        scores.append(score)
    min_width = 650 if cols >= 8 else 800 if cols >= 6 else 1000 if cols >= 4 else 1250
    widths = [max(min_width, int(CONTENT_DXA * s / sum(scores))) for s in scores]
    diff = CONTENT_DXA - sum(widths)
    if diff != 0:
        order = sorted(range(cols), key=lambda c: scores[c], reverse=True)
        step = 1 if diff > 0 else -1
        remaining = abs(diff)
        idx = 0
        while remaining:
            c = order[idx % len(order)]
            if step > 0 or widths[c] > min_width:
                widths[c] += step
                remaining -= 1
            idx += 1
    return widths


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    font_size = 7.1 if cols >= 8 else 7.6 if cols >= 7 else 8.0 if cols >= 6 else 8.4 if cols >= 5 else 8.8

    for r_idx, values in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=80, start=120, bottom=80, end=120)
            text = values[c_idx] if c_idx < len(values) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            if c_idx > 0 and len(text) <= 16 and cols >= 4:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, text, base_size=font_size, color=WHITE if r_idx == 0 else INK, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, BLUE)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, compute_widths(rows))
    set_table_borders(table, color=GRID, size="5")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph(style="Code Block")
    r = p.add_run("\n".join(lines))
    set_font(r, name="Menlo", size=6.3, color=BLUE_DARK)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "4")
        node.set(qn("w:color"), GRID)
        p_bdr.append(node)
    p_pr.append(p_bdr)


def new_decimal_num_id(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_decimal_number(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Evidence Note")
    add_inline(p, text, base_size=9.25, color=BLUE_DARK)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), BLUE_LIGHT)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def build() -> None:
    source_text = SOURCE.read_text(encoding="utf-8")
    lines = source_text.splitlines()
    doc = Document()
    configure_styles(doc)
    configure_page(doc)

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    add_cover(doc)
    add_toc(doc)

    i = 1  # Skip Markdown title; represented on cover.
    code_mode = False
    code_lines: list[str] = []
    first_h1 = True
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if code_mode:
                add_code_block(doc, code_lines)
                code_lines = []
                code_mode = False
            else:
                code_mode = True
            i += 1
            continue
        if code_mode:
            code_lines.append(line)
            i += 1
            continue

        if not stripped or stripped == "---":
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].lstrip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and first_h1:
                first_h1 = False
                style = doc.styles["Heading 1"]
                old = style.paragraph_format.page_break_before
                style.paragraph_format.page_break_before = False
                p = doc.add_paragraph(style="Heading 1")
                add_inline(p, text, base_size=16, color=BLUE, bold=True)
                style.paragraph_format.page_break_before = old
            else:
                p = doc.add_paragraph(style=f"Heading {min(level, 4)}")
                size = {1: 16, 2: 13, 3: 12, 4: 10.5}[min(level, 4)]
                color = BLUE if level <= 2 else BLUE_DARK
                add_inline(p, text, base_size=size, color=color, bold=True)
            i += 1
            continue

        if stripped.startswith("> "):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            add_note(doc, " ".join(quote_lines))
            continue

        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bullet.group(1), base_size=10.25, color=INK)
            i += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered:
            num_id = new_decimal_num_id(doc)
            while i < len(lines):
                numbered_item = re.match(r"^\d+\.\s+(.*)$", lines[i].strip())
                if not numbered_item:
                    break
                p = doc.add_paragraph(style="List Number")
                apply_decimal_number(p, num_id)
                add_inline(p, numbered_item.group(1), base_size=10.25, color=INK)
                i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt == "---" or nxt.startswith("#") or nxt.startswith("```") or nxt.startswith("|") or nxt.startswith(">") or re.match(r"^[-*]\s+", nxt) or re.match(r"^\d+\.\s+", nxt):
                break
            paragraph_lines.append(nxt)
            i += 1
        text = " ".join(paragraph_lines).replace("  ", " ")
        p = doc.add_paragraph()
        add_inline(p, text, base_size=10.25, color=INK)

    if code_mode and code_lines:
        add_code_block(doc, code_lines)

    props = doc.core_properties
    props.title = "SinhalaJournal-LLM / SINAI Repository Research Audit"
    props.subject = "Complete factual research audit for R26-SE-037"
    props.author = "R26-SE-037 Research Project"
    props.keywords = "SinhalaJournal-LLM, SINAI, SinLLaMA, research audit, Sinhala NLP"
    props.comments = "Generated from fullSummary.md; repository evidence supersedes stale README claims."

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    print(f"Paragraphs: {len(doc.paragraphs)}; tables: {len(doc.tables)}")


if __name__ == "__main__":
    build()
